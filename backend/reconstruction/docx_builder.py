"""
DOCX Builder — Main document reconstruction engine.

Orchestrates per-page reconstruction, applying styles, spacing, images,
tables, and equations to produce an editable DOCX file that faithfully
reconstructs the original document's layout.

Key improvements:
- Uses Docling/Marker text when available (higher quality than raw OCR)
- Adds page numbers in footer
- Preserves line breaks and paragraph spacing
- Proper Bangla font handling
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.config import get_settings
from reconstruction.style_manager import StyleManager
from reconstruction.page_builder import PageBuilder
from reconstruction.table_builder import TableBuilder
from reconstruction.equation_builder import EquationBuilder
from reconstruction.image_inserter import ImageInserter

logger = logging.getLogger(__name__)
settings = get_settings()


class DocxBuilder:
    """
    Main DOCX reconstruction engine.

    Takes the processed page data (layout, OCR, tables, equations) and
    produces a fully editable DOCX file preserving the original layout.
    """

    def __init__(self):
        self._doc: DocxDocument | None = None
        self._style_manager = StyleManager()
        self._page_builder = PageBuilder()
        self._table_builder = TableBuilder()
        self._equation_builder = EquationBuilder()
        self._image_inserter = ImageInserter()

    def build(
        self,
        pages_data: list[dict],
        output_path: str,
        page_width_inches: float = 8.27,  # A4
        page_height_inches: float = 11.69,  # A4
        margin_inches: float = 1.0,
        understanding_text: str = "",
    ) -> str:
        """
        Build a DOCX document from processed page data.

        Args:
            pages_data: List of per-page dicts with layout, OCR, tables, equations
            output_path: Where to save the generated DOCX
            page_width_inches: Page width (default A4)
            page_height_inches: Page height (default A4)
            margin_inches: Page margins
            understanding_text: High-quality text from Docling/Marker (preferred source)

        Returns:
            Path to the generated DOCX file
        """
        self._doc = DocxDocument()

        # Configure page layout
        self._setup_page(page_width_inches, page_height_inches, margin_inches)

        # Apply base styles
        self._style_manager.apply_default_styles(self._doc)

        # Add page numbers
        self._add_page_numbers()

        # Force page-by-page OCR-based reconstruction to preserve 1-to-1 page alignment,
        # layouts, tables, and native equations.
        logger.info("Using OCR-based page-by-page reconstruction")
        for page_idx, page_data in enumerate(pages_data):
            logger.info(f"Building page {page_idx + 1}/{len(pages_data)}")
            if page_idx > 0:
                self._doc.add_page_break()
            self._build_page(page_data)

        # Save
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_file))

        logger.info(f"DOCX saved to {output_file}")
        return str(output_file)

    def _setup_page(
        self,
        width_inches: float,
        height_inches: float,
        margin_inches: float,
    ) -> None:
        """Configure page dimensions and margins."""
        section = self._doc.sections[0]
        section.page_width = Inches(width_inches)
        section.page_height = Inches(height_inches)
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

    def _add_page_numbers(self) -> None:
        """Add page numbers to the document footer."""
        try:
            section = self._doc.sections[0]
            footer = section.footer
            footer.is_linked_to_previous = False

            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

            # Add page number field using XML
            run = para.add_run()
            run.font.size = Pt(9)
            run.font.name = "Arial"

            # PAGE field
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fldChar1)

            instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run._r.append(instrText)

            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run._r.append(fldChar2)

        except Exception as e:
            logger.warning(f"Could not add page numbers: {e}")

    def _build_page(self, page_data: dict) -> None:
        """
        Reconstruct a single page from its processed data.

        Expected page_data structure:
        {
            'page_number': int,
            'layout': {'elements': [...]},
            'ocr': {'lines': [...], 'full_text': str},
            'tables': [...],
            'equations': [...],
            'image_path': str,
        }
        """
        layout = page_data.get("layout", {})
        ocr_data = page_data.get("ocr", {})
        tables = page_data.get("tables", [])
        equations = page_data.get("equations", [])
        image_path = page_data.get("image_path", "")

        elements = layout.get("elements", [])
        ocr_lines = ocr_data.get("lines", [])

        if not elements:
            # Fallback: add all OCR text preserving line breaks
            full_text = ocr_data.get("full_text", "")
            if full_text:
                for line in full_text.split("\n"):
                    line = line.strip()
                    if line:
                        para = self._doc.add_paragraph()
                        run = para.add_run(line)
                        run.font.name = "Noto Sans Bengali"
                        run.font.size = Pt(11)
                        para.paragraph_format.space_after = Pt(4)
                        para.paragraph_format.line_spacing = 1.15
            return

        # Pre-assign OCR lines to layout elements to prevent overlap text duplication
        element_text_mapping = {i: [] for i in range(len(elements))}
        assigned_line_indices = set()

        # Sort elements by bounding box area in ascending order (smallest/most specific first)
        indexed_elements = list(enumerate(elements))
        indexed_elements.sort(key=lambda item: (item[1].get("bbox", [0,0,0,0])[2] - item[1].get("bbox", [0,0,0,0])[0]) * 
                                              (item[1].get("bbox", [0,0,0,0])[3] - item[1].get("bbox", [0,0,0,0])[1]))

        for el_idx, element in indexed_elements:
            bbox = element.get("bbox", [0, 0, 0, 0])
            for line_idx, line in enumerate(ocr_lines):
                if line_idx in assigned_line_indices:
                    continue
                line_bbox = line.get("bbox", [0, 0, 0, 0])
                if self._bboxes_overlap(bbox, line_bbox):
                    element_text_mapping[el_idx].append(line)
                    assigned_line_indices.add(line_idx)

        # Process elements in reading order
        table_idx = 0
        equation_idx = 0

        for el_idx, element in sorted(list(enumerate(elements)), key=lambda e: e[1].get("reading_order", 0)):
            element_type = element.get("type", "unknown")
            assigned_lines = element_text_mapping[el_idx]

            try:
                if element_type == "title":
                    self._add_title(element, assigned_lines)
                elif element_type == "subtitle":
                    self._add_subtitle(element, assigned_lines)
                elif element_type in ("paragraph", "unknown"):
                    self._add_paragraph(element, assigned_lines)
                elif element_type == "table" and table_idx < len(tables):
                    self._table_builder.add_table(self._doc, tables[table_idx])
                    table_idx += 1
                elif element_type == "equation" and equation_idx < len(equations):
                    self._equation_builder.add_equation(self._doc, equations[equation_idx])
                    equation_idx += 1
                elif element_type in ("image", "figure"):
                    self._image_inserter.add_image(self._doc, element, image_path)
                elif element_type == "caption":
                    self._add_caption(element, assigned_lines)
                elif element_type == "header":
                    pass  # Headers handled separately
                elif element_type == "footer":
                    pass  # Footers handled separately
                elif element_type == "list":
                    self._add_list(element, assigned_lines)
                elif element_type == "exercise":
                    self._add_exercise(element, assigned_lines)
                else:
                    # Default: treat as paragraph
                    self._add_paragraph(element, assigned_lines)
            except Exception as e:
                logger.warning(f"Error processing {element_type}: {e}")
                # Fallback: add as plain text
                text = "\n".join(line.get("text", "") for line in assigned_lines) if assigned_lines else ""
                if text:
                    self._doc.add_paragraph(text)

        # Fallback for any unassigned text lines to ensure 100% text coverage
        unassigned_lines = []
        for line_idx, line in enumerate(ocr_lines):
            if line_idx not in assigned_line_indices:
                unassigned_lines.append(line)
        if unassigned_lines:
            logger.info(f"Page {page_data.get('page_number')}: adding {len(unassigned_lines)} unassigned OCR lines as fallback")
            para = self._doc.add_paragraph()
            text = "\n".join(line.get("text", "") for line in unassigned_lines)
            run = para.add_run(text)
            run.font.name = "Noto Sans Bengali"
            run.font.size = Pt(11)

    def _bboxes_overlap(self, bbox1: list, bbox2: list) -> bool:
        """Check if two bounding boxes overlap significantly."""
        x1 = max(bbox1[0], bbox2[0])
        y1 = max(bbox1[1], bbox2[1])
        x2 = min(bbox1[2], bbox2[2])
        y2 = min(bbox1[3], bbox2[3])

        if x1 >= x2 or y1 >= y2:
            return False

        intersection = (x2 - x1) * (y2 - y1)
        area2 = (bbox2[2] - bbox2[0]) * (bbox2[3] - bbox2[1])

        return intersection / max(area2, 1) > 0.3

    def _add_title(self, element: dict, assigned_lines: list[dict]) -> None:
        text = " ".join(line.get("text", "") for line in assigned_lines) if assigned_lines else ""
        if text:
            para = self._doc.add_heading(text, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_subtitle(self, element: dict, assigned_lines: list[dict]) -> None:
        text = " ".join(line.get("text", "") for line in assigned_lines) if assigned_lines else ""
        if text:
            self._doc.add_heading(text, level=2)

    def _add_paragraph(self, element: dict, assigned_lines: list[dict]) -> None:
        text = "\n".join(line.get("text", "") for line in assigned_lines) if assigned_lines else ""
        if text:
            para = self._doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = "Noto Sans Bengali"
            run.font.size = Pt(11)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.line_spacing = 1.15

    def _add_caption(self, element: dict, assigned_lines: list[dict]) -> None:
        text = " ".join(line.get("text", "") for line in assigned_lines) if assigned_lines else ""
        if text:
            para = self._doc.add_paragraph()
            para.style = "Caption" if "Caption" in self._doc.styles else para.style
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = "Noto Sans Bengali"

    def _add_list(self, element: dict, assigned_lines: list[dict]) -> None:
        if assigned_lines:
            for line in assigned_lines:
                item = line.get("text", "").strip()
                if item:
                    para = self._doc.add_paragraph(item, style="List Bullet")
                    for run in para.runs:
                        run.font.name = "Noto Sans Bengali"

    def _add_exercise(self, element: dict, assigned_lines: list[dict]) -> None:
        text = "\n".join(line.get("text", "") for line in assigned_lines) if assigned_lines else ""
        if text:
            para = self._doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Noto Sans Bengali"
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
