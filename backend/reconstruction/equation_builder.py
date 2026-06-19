"""
Equation Builder — Insert equations into DOCX.

Converts LaTeX equations to Office Math Markup Language (OMML) for
native Word equation support. Falls back to equation images.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class EquationBuilder:
    """Inserts mathematical equations into DOCX documents."""

    def add_equation(self, doc: DocxDocument, equation_data: dict) -> None:
        """
        Add an equation to the document.

        Expected equation_data:
        {
            'latex': str,
            'type': 'display' | 'inline',
            'image_path': str (optional fallback),
            'confidence': float,
        }
        """
        latex = equation_data.get("latex", "")
        eq_type = equation_data.get("type", "display")
        image_path = equation_data.get("image_path", "")

        if not latex and not image_path:
            return

        # Try to insert as OMML (native Word equation)
        if latex:
            try:
                self._add_omml_equation(doc, latex, eq_type)
                return
            except Exception as e:
                logger.warning(f"OMML insertion failed, falling back to image: {e}")

        # Fallback: try math2docx
        if latex:
            try:
                self._add_math2docx(doc, latex)
                return
            except Exception as e:
                logger.warning(f"math2docx failed: {e}")

        # Final fallback: insert equation as image
        if image_path and Path(image_path).exists():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Inches(4))

    def _add_omml_equation(self, doc: DocxDocument, latex: str, eq_type: str) -> None:
        """
        Convert LaTeX to OMML and insert into the document natively.
        Uses latex2mathml and lxml for conversion.
        """
        import latex2mathml.converter
        from lxml import etree
        from docx.oxml import parse_xml

        # Convert LaTeX to MathML
        mathml_str = latex2mathml.converter.convert(latex)

        # Load XSLT stylesheet from backend/reconstruction/MML2OMML.XSL
        xslt_path = Path(__file__).parent / "MML2OMML.XSL"
        if not xslt_path.exists():
            xslt_path = Path("reconstruction/MML2OMML.XSL")

        if not xslt_path.exists():
            raise FileNotFoundError(f"MML2OMML.XSL not found at {xslt_path.absolute()}")

        xslt_tree = etree.parse(str(xslt_path))
        transform = etree.XSLT(xslt_tree)

        # Transform MathML to OMML
        mathml_tree = etree.fromstring(mathml_str)
        omml_tree = transform(mathml_tree)

        # Get XML string of the generated OMML element
        omml_xml = etree.tostring(omml_tree, encoding="utf-8").decode("utf-8")

        # Parse XML element
        omml_element = parse_xml(omml_xml)

        para = doc.add_paragraph()
        if eq_type == "display":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Append the native OMML element to the paragraph
        para._element.append(omml_element)

    def _add_math2docx(self, doc: DocxDocument, latex: str) -> None:
        """Use math2docx library for equation insertion."""
        try:
            from math2docx import add_math
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_math(para, latex)
        except ImportError:
            raise
