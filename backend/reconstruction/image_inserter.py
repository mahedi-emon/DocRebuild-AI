"""
Image Inserter — Places images in DOCX documents.

Crops image regions from original pages and inserts them with captions.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.image_utils import crop_region, load_image, save_image

logger = logging.getLogger(__name__)


class ImageInserter:
    """Inserts images into DOCX documents."""

    def add_image(
        self,
        doc: DocxDocument,
        element: dict,
        source_image_path: str,
        max_width_inches: float = 5.0,
    ) -> None:
        """
        Add an image from a layout element to the document.

        Crops the image region from the source page image and inserts it.
        """
        bbox = element.get("bbox", [])
        caption_text = element.get("caption", "")

        if not source_image_path or not Path(source_image_path).exists():
            return

        try:
            # Crop the image region
            source = load_image(source_image_path)
            if bbox and len(bbox) == 4:
                region = crop_region(
                    source,
                    (int(bbox[0]), int(bbox[1]), int(bbox[2]), int(bbox[3])),
                    padding=5,
                )
            else:
                region = source

            # Save cropped region to temp file
            temp_path = Path(source_image_path).parent / f"_crop_{id(element)}.png"
            save_image(region, str(temp_path))

            # Calculate appropriate width
            h, w = region.shape[:2]
            aspect = w / h if h > 0 else 1
            width = min(max_width_inches, aspect * 3.0)

            # Insert image
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(temp_path), width=Inches(width))

            # Add caption if available
            if caption_text:
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap_para.add_run(caption_text)
                run.italic = True
                run.font.size = Pt(9)

        except Exception as e:
            logger.warning(f"Failed to insert image: {e}")
