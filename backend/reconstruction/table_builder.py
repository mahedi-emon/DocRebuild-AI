"""
Table Builder — Reconstructs extracted tables as editable DOCX tables.
"""

from __future__ import annotations

import logging

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


class TableBuilder:
    """Converts extracted table data into DOCX tables."""

    def add_table(self, doc: DocxDocument, table_data: dict) -> None:
        """
        Add a table to the document.

        Expected table_data format:
        {
            'rows': int,
            'columns': int,
            'cells': [[str, ...], ...],  # 2D array of cell text
            'headers': [str, ...],  # Optional header row
        }
        """
        rows = table_data.get("rows", 0)
        cols = table_data.get("columns", 0)
        cells = table_data.get("cells", [])
        headers = table_data.get("headers", [])

        if not rows or not cols:
            # Try to infer from cells
            if cells:
                rows = len(cells)
                cols = max(len(row) for row in cells) if cells else 0

        if not rows or not cols:
            logger.warning("Table has no rows/columns, skipping")
            return

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill cells
        for r_idx, row_data in enumerate(cells[:rows]):
            for c_idx, cell_text in enumerate(row_data[:cols]):
                cell = table.cell(r_idx, c_idx)
                cell.text = str(cell_text) if cell_text else ""
                # Style cell text
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(10)

        # Bold header row if present
        if headers or (cells and len(cells) > 0):
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

        # Add spacing after table
        doc.add_paragraph()
